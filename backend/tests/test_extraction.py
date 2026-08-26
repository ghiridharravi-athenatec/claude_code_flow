"""Unit tests for the per-type extraction modules."""

from __future__ import annotations

import io

import pytest
from docx import Document

from extraction.docx_extractor import extract_docx_text
from extraction.errors import CorruptedFileError, EmptyContentError
from extraction.pdf_extractor import extract_pdf_text
from extraction.txt_extractor import extract_txt_text


# --- txt_extractor ---------------------------------------------------------


def test_txt_extracts_utf8():
    text = extract_txt_text("Chief complaint: chest pain, well documented.".encode("utf-8"))
    assert "chest pain" in text


def test_txt_falls_back_to_latin1():
    raw = "Café notes: patient reports discomfort in the chest area.".encode("latin-1")
    text = extract_txt_text(raw)
    assert "Café" in text or "Caf" in text


def test_txt_utf8_decode_failure_falls_back_to_latin1_not_decode_error():
    # Every byte value is valid latin-1, so the utf-8-fails/latin-1-fails DecodeError
    # branch is unreachable for a bytes input -- this documents that, rather than
    # asserting a case that cannot occur.
    raw = b"\xff\xfe invalid utf-8 but valid latin-1 chest pain notes"
    text = extract_txt_text(raw)
    assert "chest pain" in text


def test_txt_raises_empty_content_error_on_near_empty_text():
    with pytest.raises(EmptyContentError):
        extract_txt_text(b"   \n\t  ")


# --- docx_extractor ---------------------------------------------------------


def _build_docx_bytes(paragraphs: list[str]) -> bytes:
    document = Document()
    for paragraph in paragraphs:
        document.add_paragraph(paragraph)
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def test_docx_extracts_paragraph_text():
    file_bytes = _build_docx_bytes(["Chief complaint: shortness of breath on exertion."])
    text = extract_docx_text(file_bytes)
    assert "shortness of breath" in text


def test_docx_raises_empty_content_error_on_blank_document():
    file_bytes = _build_docx_bytes([])
    with pytest.raises(EmptyContentError):
        extract_docx_text(file_bytes)


def test_docx_raises_corrupted_error_on_garbage_bytes():
    with pytest.raises(CorruptedFileError):
        extract_docx_text(b"not a real docx file")


def test_docx_extracts_table_cell_text():
    document = Document()
    table = document.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "Medication"
    table.rows[0].cells[1].text = "Lisinopril 10mg daily"
    buffer = io.BytesIO()
    document.save(buffer)

    text = extract_docx_text(buffer.getvalue())

    assert "Lisinopril" in text


# --- pdf_extractor ---------------------------------------------------------


def test_pdf_raises_corrupted_error_on_garbage_bytes():
    with pytest.raises(CorruptedFileError):
        extract_pdf_text(b"not a real pdf file")


def test_pdf_raises_empty_content_error_on_blank_page():
    from pypdf import PdfWriter

    writer = PdfWriter()
    writer.add_blank_page(width=200, height=200)
    buffer = io.BytesIO()
    writer.write(buffer)

    with pytest.raises(EmptyContentError):
        extract_pdf_text(buffer.getvalue())
