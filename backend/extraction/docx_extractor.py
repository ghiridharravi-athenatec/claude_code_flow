"""Text extraction for DOCX uploads."""

from __future__ import annotations

import io

from docx import Document
from docx.opc.exceptions import PackageNotFoundError

from extraction.errors import CorruptedFileError, EmptyContentError

MIN_NON_WHITESPACE_CHARS = 20


def extract_docx_text(file_bytes: bytes) -> str:
    """Concatenate paragraph text plus table cell text (row-major, tab-joined)."""
    try:
        document = Document(io.BytesIO(file_bytes))
    except PackageNotFoundError as exc:
        raise CorruptedFileError() from exc
    except Exception as exc:
        raise CorruptedFileError() from exc

    lines = [paragraph.text for paragraph in document.paragraphs]

    for table in document.tables:
        for row in table.rows:
            lines.append("\t".join(cell.text for cell in row.cells))

    text = "\n".join(lines)

    if len("".join(text.split())) < MIN_NON_WHITESPACE_CHARS:
        raise EmptyContentError()

    return text
